import os
import csv
import logging
from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)

def extract_pdf(file_path):
    """
    Extracts text from a PDF file using Docling (via Nova Agent), with a fallback to pypdf.
    """
    # 1. Try deep parsing via Nova Agent (Docling)
    try:
        from app.nova.services import NovaAgentService
        logger.info("[User Services] Delegando extração de PDF para Nova Agent...")
        agent = NovaAgentService()
        markdown_text = agent.process_document(file_path)
        if markdown_text:
            return markdown_text
    except Exception as e:
        logger.warning(f"[User Services] Nova Agent falhou ou Docling indisponível. Utilizando fallback (pypdf): {e}")

    # 2. Robust fallback using pypdf
    import pypdf
    try:
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"## --- Página {i + 1} ---\n\n{text}")
        return "\n\n".join(text_parts) if text_parts else "PDF vazio ou sem texto extraível."
    except Exception as e:
        logger.error(f"Erro ao extrair PDF via pypdf: {e}")
        return f"Falha na extração de texto do PDF: {str(e)}"

def extract_excel(file_path):
    """
    Extracts data from an Excel workbook and converts sheets into beautiful Markdown tables.
    """
    import openpyxl
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        md_parts = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            md_parts.append(f"### 📊 Aba: {sheet_name}\n")
            
            # Extract rows
            rows = list(sheet.iter_rows(values_only=True))
            if not rows or all(all(val is None for val in r) for r in rows):
                md_parts.append("*Esta aba está vazia.*\n")
                continue
                
            # Filter empty trailing rows/columns
            non_empty_rows = [r for r in rows if not all(val is None for val in r)]
            if not non_empty_rows:
                md_parts.append("*Esta aba está vazia.*\n")
                continue
                
            # Find max active columns
            max_cols = max(len(r) for r in non_empty_rows)
            
            # Format header row
            headers = [str(val) if val is not None else "" for val in non_empty_rows[0]]
            # Fill headers if row is shorter
            while len(headers) < max_cols:
                headers.append("")
            
            header_line = "| " + " | ".join(headers) + " |"
            sep_line = "| " + " | ".join(["---"] * max_cols) + " |"
            md_parts.append(header_line)
            md_parts.append(sep_line)
            
            # Format data rows
            for row in non_empty_rows[1:]:
                row_vals = [str(val) if val is not None else "" for val in row]
                while len(row_vals) < max_cols:
                    row_vals.append("")
                # Escape pipe symbols in cell values to not break markdown table
                row_vals = [v.replace("|", "\\|").replace("\n", " ") for v in row_vals]
                md_parts.append("| " + " | ".join(row_vals) + " |")
            
            md_parts.append("\n") # Blank line after table
            
        return "\n".join(md_parts) if md_parts else "Planilha Excel vazia."
    except Exception as e:
        logger.error(f"Erro ao extrair Excel: {e}")
        return f"Falha na extração de dados do Excel: {str(e)}"

def extract_csv(file_path):
    """
    Extracts text from CSV and formats it into a beautiful Markdown table.
    """
    try:
        md_parts = []
        with open(file_path, mode='r', encoding='utf-8-sig', errors='replace') as f:
            # Try to auto-detect delimiter
            sample = f.read(2048)
            f.seek(0)
            delimiter = ','
            if ';' in sample:
                delimiter = ';'
            elif '\t' in sample:
                delimiter = '\t'
                
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
            
        if not rows:
            return "CSV vazio."
            
        max_cols = max(len(r) for r in rows)
        
        # Header
        headers = [str(val) if val else f"Coluna {i+1}" for i, val in enumerate(rows[0])]
        while len(headers) < max_cols:
            headers.append("")
        header_line = "| " + " | ".join(headers) + " |"
        sep_line = "| " + " | ".join(["---"] * max_cols) + " |"
        md_parts.append(header_line)
        md_parts.append(sep_line)
        
        # Rows
        for row in rows[1:]:
            row_vals = [str(val) if val is not None else "" for val in row]
            while len(row_vals) < max_cols:
                row_vals.append("")
            row_vals = [v.replace("|", "\\|").replace("\n", " ") for v in row_vals]
            md_parts.append("| " + " | ".join(row_vals) + " |")
            
        return "\n".join(md_parts)
    except Exception as e:
        logger.error(f"Erro ao extrair CSV: {e}")
        return f"Falha na extração de dados do CSV: {str(e)}"

def extract_docx(file_path):
    """
    Extracts text and headings from a Word file (.docx) using Docling (via Nova Agent), with a fallback to python-docx.
    """
    # 1. Try deep parsing via Nova Agent (Docling)
    try:
        from app.nova.services import NovaAgentService
        logger.info("[User Services] Delegando extração de Word para Nova Agent...")
        agent = NovaAgentService()
        markdown_text = agent.process_document(file_path)
        if markdown_text:
            return markdown_text
    except Exception as e:
        logger.warning(f"[User Services] Nova Agent falhou ou Docling indisponível. Utilizando fallback (python-docx): {e}")

    # 2. Robust fallback using python-docx
    import docx
    try:
        doc = docx.Document(file_path)
        md_parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Map word headings to markdown headings
            if p.style.name.startswith('Heading 1'):
                md_parts.append(f"# {text}\n")
            elif p.style.name.startswith('Heading 2'):
                md_parts.append(f"## {text}\n")
            elif p.style.name.startswith('Heading 3'):
                md_parts.append(f"### {text}\n")
            elif p.style.name.startswith('List'):
                md_parts.append(f"* {text}")
            else:
                md_parts.append(text)
                
        # Also extract tables from docx
        for table in doc.tables:
            md_parts.append("\n### 📋 Tabela Extraída do Word\n")
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ").replace("|", "\\|") for cell in row.cells]
                table_rows.append(cells)
            
            if not table_rows:
                continue
                
            max_cols = max(len(r) for r in table_rows)
            # Header
            headers = table_rows[0]
            while len(headers) < max_cols:
                headers.append("")
            md_parts.append("| " + " | ".join(headers) + " |")
            md_parts.append("| " + " | ".join(["---"] * max_cols) + " |")
            # Data
            for row in table_rows[1:]:
                while len(row) < max_cols:
                    row.append("")
                md_parts.append("| " + " | ".join(row) + " |")
            md_parts.append("\n")

        return "\n\n".join(md_parts) if md_parts else "Word vazio ou sem texto extraível."
    except Exception as e:
        logger.error(f"Erro ao extrair Word: {e}")
        return f"Falha na extração de dados do Word: {str(e)}"

def extract_markdown(file_path):
    """
    Reads plain text/markdown directly.
    """
    try:
        with open(file_path, mode='r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        return content if content.strip() else "Arquivo Markdown original está vazio."
    except Exception as e:
        logger.error(f"Erro ao ler Markdown: {e}")
        return f"Falha na leitura do arquivo Markdown: {str(e)}"

def process_document_file(documento):
    """
    Core function that extracts text, saves the Markdown representation, 
    and generates the structured AI analysis (analise_ia).
    """
    if not documento.arquivo:
        return
        
    file_path = documento.arquivo.path
    ext = documento.extensao
    
    logger.info(f"Iniciando processamento do arquivo: {file_path} com extensão {ext}")
    
    # 1. Extract content based on extension
    if ext == '.pdf':
        content = extract_pdf(file_path)
        format_name = "PDF (Portable Document Format)"
        icon = "📕"
    elif ext in ['.xlsx', '.xls']:
        content = extract_excel(file_path)
        format_name = "Planilha Excel"
        icon = "📗"
    elif ext == '.csv':
        content = extract_csv(file_path)
        format_name = "Arquivo de Valores Separados por Vírgula (CSV)"
        icon = "📊"
    elif ext == '.docx':
        content = extract_docx(file_path)
        format_name = "Documento do Word"
        icon = "📘"
    elif ext == '.md':
        content = extract_markdown(file_path)
        format_name = "Documento Markdown"
        icon = "📝"
    else:
        content = f"Formato não suportado ativamente para extração profunda: {ext}"
        format_name = "Outro / Desconhecido"
        icon = "📄"
        
    # 2. Save Markdown representation (.md file)
    md_filename = os.path.basename(file_path)
    # Remove original extension and add _processado.md
    name_parts = os.path.splitext(md_filename)
    md_filename = f"{name_parts[0]}_processado.md"
    
    # Save the processed markdown file
    documento.arquivo_markdown.save(
        md_filename,
        ContentFile(content.encode('utf-8')),
        save=False
    )
    
    # 3. Generate structured AI Analysis (analise_ia)
    file_size_kb = os.path.getsize(file_path) / 1024
    
    summary_card = f"""# {icon} Relatório de Processamento do Documento

> [!NOTE]
> **Status:** Processamento concluído com sucesso via Padrão de Projeto Observer (Django Signals).
> 
> * **Nome do Documento:** `{documento.nome}`
> * **Formato Original:** {format_name} (`{ext}`)
> * **Tamanho do Arquivo:** `{file_size_kb:.2f} KB`
> * **Caminho Físico:** `{os.path.basename(file_path)}`

---

## 🔍 Visão Geral & Estrutura do Documento

Este documento foi analisado e estruturado automaticamente por nossos agentes de processamento. A estrutura física foi catalogada e disponibilizada para consumo imediato do agente de inteligência artificial.

### 📋 Principais Seções Extraídas

"""
    
    # Generate appropriate layout based on content size and type
    if ext in ['.xlsx', '.xls', '.csv']:
        summary_card += f"### 📊 Tabelas de Dados Encontradas\n\nPlanilha contendo dados estruturados em linhas e colunas. Abaixo estão as tabelas extraídas diretamente:\n\n{content}\n"
    else:
        # Create a preview of content for other text documents
        lines = content.split('\n')
        preview_lines = []
        char_count = 0
        for line in lines:
            if char_count < 1500:
                preview_lines.append(line)
                char_count += len(line)
            else:
                preview_lines.append("\n\n*... [Conteúdo restante truncado para visualização rápida. O arquivo Markdown completo está disponível para download] ...*")
                break
        preview_content = "\n".join(preview_lines)
        summary_card += f"### 📝 Pré-visualização do Conteúdo Extraído\n\n{preview_content}\n"
        
    summary_card += f"""
---

## 🤖 Recomendações para Agentes de IA

Este documento possui dados de alto valor. Para obter os melhores resultados com seu agente de IA:
1. **Perguntas Estruturadas:** Faça consultas focadas na análise ou tópicos específicos do documento.
2. **Consultas de Tabela:** O agente de IA pode ler diretamente as tabelas Markdown extraídas acima para realizar cálculos ou criar resumos consolidados.
3. **Download Completo:** Se necessário, você pode baixar o arquivo markdown gerado (`{md_filename}`) ou o arquivo original clicando nos botões de ação na barra superior.
"""
    
    documento.analise_ia = summary_card
    documento.save()
    logger.info(f"Processamento concluído com sucesso para o documento: {documento.nome}")
